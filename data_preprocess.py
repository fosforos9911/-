import re
import logging
from pathlib import Path
from typing import Dict, List, Union, Optional, Any

import pandas as pd
import pdfplumber
from docx import Document
import chardet
import time

try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False
    # 定义一个简单的替代函数
    def tqdm(iterable, desc="", **kwargs):
        print(f"\n{desc}")
        for item in iterable:
            print(".", end="", flush=True)
            yield item
        print(" 完成!")



def process_all_files_in_data():
    """处理data文件夹中的所有支持的文件"""
    
    # 获取当前文件所在目录的上一级，然后加上data文件夹
    current_dir = Path(__file__).parent
    data_dir = current_dir / "data"
    
    # 检查data文件夹是否存在
    if not data_dir.exists():
        print(f"警告：data文件夹不存在 - {data_dir}")
        return []
    
    # 支持的文件扩展名
    supported_extensions = ['.pdf', '.docx', '.doc', '.csv', '.txt']
    
    results = []
    
    # 获取所有支持的文件
    files_to_process = [f for f in data_dir.iterdir() 
                       if f.is_file() and f.suffix.lower() in supported_extensions]
    
    if not files_to_process:
        print("data文件夹中没有支持的文件")
        return []
    
    print(f"找到 {len(files_to_process)} 个文件待处理")
    
    # 遍历data文件夹中的所有文件
    for file_path in tqdm(files_to_process, desc="处理文件"):
        try:
            print(f"\n正在处理: {file_path.name}")
            
            # 调用提取函数
            result = extract_and_convert(str(file_path))
            results.append({
                'file_name': file_path.name,
                'file_path': str(file_path),
                'data': result
            })
            
            print(f"✓ 处理完成: {file_path.name}")
            
        except Exception as e:
            print(f"✗ 处理失败: {file_path.name}, 错误: {str(e)}")
    
    summary_msg = f"\n总共处理了 {len(results)} 个文件"
    print(summary_msg)
    
    return results


def read_pdf_file(file_path: str) -> Dict[str, Union[str, List[List[str]]]]:
    """
    读取PDF文件，提取文本和表格，保留表格结构
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        包含文本和表格的字典: {
            'text': 提取的纯文本,
            'tables': 表格列表，每个表格为二维列表,
            'pages': 分页内容列表
        }
    """
    result = {
        'text': '',
        'tables': [],
        'pages': []
    }
    
    try:
        # 使用pdfplumber提取（更适合表格识别）
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # 提取页面文本
                page_text = page.extract_text() or ''
                
                # 提取页面中的表格
                tables = page.extract_tables()
                
                # 处理每个表格，保留行列关系
                for table_idx, table in enumerate(tables):
                    if table:  # 确保表格不为空
                        # 清理表格数据：去除None值，保留行列结构
                        clean_table = []
                        for row in table:
                            clean_row = [cell if cell is not None else '' for cell in row]
                            if any(clean_row):  # 只保留非空行
                                clean_table.append(clean_row)
                        
                        if clean_table:  # 确保清理后的表格不为空
                            result['tables'].append({
                                'page': page_num + 1,
                                'index': len(result['tables']) + 1,
                                'data': clean_table,
                                'shape': (len(clean_table), len(clean_table[0]) if clean_table else 0)
                            })
                            
                            # 将表格以Markdown格式添加到文本中，保留结构
                            page_text += f"\n\n[表格 {len(result['tables'])} 在第{page_num + 1}页]\n"
                            for row in clean_table[:3]:  # 只显示前3行作为预览
                                page_text += ' | '.join(str(cell)[:20] for cell in row) + '\n'
                
                result['pages'].append({
                    'page_num': page_num + 1,
                    'text': page_text,
                    'table_count': len(tables)
                })
                
                result['text'] += page_text + '\n\n'
    
    except Exception as e:
        print(f"PDF读取失败: {file_path}, 错误: {str(e)}")
        raise
    
    return result


def read_doc_file(file_path: str) -> Dict[str, Union[str, List[List[str]]]]:
    """
    读取Word文档，提取文本和表格，保留表格结构
    
    Args:
        file_path: Word文档路径
        
    Returns:
        包含文本和表格的字典: {
            'text': 提取的纯文本,
            'tables': 表格列表，每个表格为二维列表,
            'paragraphs': 段落列表
        }
    """
    result = {
        'text': '',
        'tables': [],
        'paragraphs': []
    }
    
    try:
        # 检查文件是否存在
        if not Path(file_path).exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 打开Word文档
        doc = Document(file_path)
        
        # 提取所有段落
        for para in doc.paragraphs:
            if para.text.strip():  # 只保留非空段落
                result['paragraphs'].append(para.text)
                result['text'] += para.text + '\n'
        
        # 提取所有表格
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            
            # 遍历表格的所有行
            for row in table.rows:
                row_data = []
                # 遍历行中的所有单元格
                for cell in row.cells:
                    # 合并单元格中的文本（可能包含多个段落）
                    cell_text = ' '.join([p.text for p in cell.paragraphs])
                    row_data.append(cell_text.strip())
                
                # 只添加非空行
                if any(cell.strip() for cell in row_data):
                    table_data.append(row_data)
            
            if table_data:  # 确保表格不为空
                result['tables'].append({
                    'index': table_idx + 1,
                    'data': table_data,
                    'shape': (len(table_data), len(table_data[0]) if table_data else 0)
                })
                
                # 将表格以结构化格式添加到文本
                result['text'] += f"\n[表格 {table_idx + 1}]\n"
                for row in table_data[:3]:  # 预览前3行
                    result['text'] += ' | '.join(str(cell)[:30] for cell in row) + '\n'
                result['text'] += '\n'
        
    except Exception as e:
        print(f"Word文档读取失败: {file_path}, 错误: {str(e)}")
        raise
    
    return result


def read_csv_file(file_path: str, encoding: Optional[str] = None) -> Dict[str, Any]:
    """
    读取CSV文件，自动检测编码，保留原始数据结构
    
    Args:
        file_path: CSV文件路径
        encoding: 指定编码，如果不指定则自动检测
        
    Returns:
        包含CSV数据的字典: {
            'data': DataFrame格式的数据,
            'headers': 列标题,
            'shape': (行数, 列数),
            'preview': 前5行数据预览,
            'dtypes': 各列数据类型
        }
    """
    result = {}
    df = None  # 初始化df
    detected_encoding = encoding
    
    try:
        # 检查文件是否存在
        if not Path(file_path).exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 如果没有指定编码，自动检测
        if encoding is None:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_result = chardet.detect(raw_data)
                detected_encoding = encoding_result['encoding'] or 'utf-8'
                confidence = encoding_result['confidence']
                logging.info(f"检测到编码: {detected_encoding} (置信度: {confidence:.2%})")
        
        # 尝试读取CSV
        df = pd.read_csv(file_path, encoding=detected_encoding)
        
        # 处理可能的编码问题
        if df.empty:
            # 如果pandas读取为空，尝试用csv模块读取
            print(f"Pandas读取为空，尝试其他方法: {file_path}")
            raise ValueError("Pandas读取为空")
        
    except Exception as e:
        print(f"CSV读取失败: {file_path}, 错误: {str(e)}")
        # 兼容其他常见编码
        common_encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1', 'cp1252', 'big5']
        for enc in common_encodings:
            if enc != detected_encoding:
                try:
                    print(f"尝试编码: {enc}")
                    return read_csv_file(file_path, encoding=enc)
                except:
                    continue
        
        # 如果所有编码失败，返回空结果
        return {
            'data': pd.DataFrame(),
            'headers': [],
            'shape': (0, 0),
            'preview': [],
            'dtypes': {},
            'encoding': detected_encoding or 'unknown',
            'error': str(e),
            'warning': '文件读取失败'
        }
    
    # 将处理逻辑移到try块外面，确保df已定义
    if df is not None and not df.empty:
        # 去除完全空的行和列
        df = df.dropna(how='all').dropna(axis=1, how='all')
        
        # 尝试转换数值类型
        for col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='ignore')
        
        # 构建结果
        result = {
            'data': df,
            'headers': df.columns.tolist(),
            'shape': (len(df), len(df.columns)),
            'preview': df.head(5).to_dict('records'),
            'dtypes': df.dtypes.astype(str).to_dict(),
            'encoding': detected_encoding,
            'missing_values': df.isnull().sum().to_dict(),
            'memory_usage': df.memory_usage(deep=True).sum() / 1024 / 1024  # MB
        }
    else:
        result = {
            'data': pd.DataFrame() if df is None else df,
            'headers': [],
            'shape': (0, 0),
            'preview': [],
            'dtypes': {},
            'encoding': detected_encoding or 'unknown',
            'warning': '文件为空'
        }
    
    return result


def extract_and_convert(file_path: str) -> Dict[str, Any]:
    """
    统一入口：根据文件类型自动提取文本和表格，并转换为标准格式
    
    Args:
        file_path: 文件路径（支持PDF、DOCX、CSV、TXT）
        
    Returns:
        标准化的数据字典：{
            'raw_text': 原始文本内容,
            'tables': 表格列表（每个表格为二维列表）,
            'metadata': 元数据（文件类型、页数等）
        }
    """
    
    # 添加文件存在性检查
    if not Path(file_path).exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    def detect_file_type(path: str) -> str:
        """检测文件类型"""
        ext = Path(path).suffix.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext in ['.docx', '.doc']:
            return 'doc'
        elif ext == '.csv':
            return 'csv'
        elif ext == '.txt':
            return 'txt'
        else:
            raise ValueError(f"不支持的文件类型: {ext}")
    
    def extract_from_pdf(path: str) -> Dict:
        """从PDF提取文本和表格"""
        text_content = ""
        tables_list = []
        total_pages = 0
        
        with pdfplumber.open(path) as pdf:
            total_pages = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages):
                # 提取文本
                page_text = page.extract_text() or ''
                text_content += f"\n--- 第 {page_num + 1} 页 ---\n{page_text}\n"
                
                # 提取表格并转换为Markdown格式嵌入文本
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables):
                    if table and len(table) > 0:
                        # 清理表格
                        clean_table = []
                        for row in table:
                            clean_row = [str(cell).strip() if cell else '' for cell in row]
                            if any(clean_row):
                                clean_table.append(clean_row)
                        
                        if clean_table:
                            tables_list.append({
                                'page': page_num + 1,
                                'index': len(tables_list) + 1,
                                'data': clean_table,
                                'shape': (len(clean_table), len(clean_table[0]) if clean_table else 0)
                            })
                            
                            # 将表格以Markdown格式添加到文本
                            text_content += f"\n[表格 {len(tables_list)} 在第{page_num + 1}页]\n"
                            header = clean_table[0]
                            text_content += '| ' + ' | '.join(header) + ' |\n'
                            text_content += '|' + '|'.join([' --- ' for _ in header]) + '|\n'
                            
                            for row in clean_table[1:6]:  # 显示前5行数据
                                # 确保行长度与表头一致
                                while len(row) < len(header):
                                    row.append('')
                                text_content += '| ' + ' | '.join(row) + ' |\n'
                            if len(clean_table) > 6:
                                text_content += '| ... |\n'
                            text_content += '\n'
        
        return {
            'raw_text': text_content,
            'tables': tables_list,
            'metadata': {'type': 'pdf', 'pages': total_pages}
        }
    
    def extract_from_doc(path: str) -> Dict:
        """从Word文档提取文本和表格"""
        doc = Document(path)
        text_content = ""
        tables_list = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_content += para.text + '\n'
        
        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = ' '.join([p.text for p in cell.paragraphs])
                    row_data.append(cell_text.strip())
                if any(row_data):
                    table_data.append(row_data)
            
            if table_data:
                tables_list.append({
                    'index': len(tables_list) + 1,
                    'data': table_data,
                    'shape': (len(table_data), len(table_data[0]) if table_data else 0)
                })
                
                text_content += f"\n[表格 {len(tables_list)}]\n"
                for row in table_data[:5]:  # 只显示前5行预览
                    text_content += ' | '.join(row) + '\n'
                if len(table_data) > 5:
                    text_content += '...\n'
                text_content += '\n'
        
        return {
            'raw_text': text_content,
            'tables': tables_list,
            'metadata': {'type': 'doc', 'paragraphs': len(doc.paragraphs)}
        }
    
    def extract_from_csv(path: str) -> Dict:
        """从CSV提取数据"""
        # 调用read_csv_file函数
        csv_result = read_csv_file(path)
        
        # 转换为文本描述
        if csv_result['shape'][0] > 0:
            text_content = f"CSV数据: {csv_result['shape'][0]}行 × {csv_result['shape'][1]}列\n"
            text_content += f"列名: {', '.join(csv_result['headers'])}\n\n"
            text_content += "前5行预览:\n"
            
            # 将预览数据转换为字符串
            df = csv_result['data']
            text_content += df.head(5).to_string() + '\n'
            
            # 添加统计信息
            text_content += f"\n内存占用: {csv_result.get('memory_usage', 0):.2f} MB\n"
            
            # 将整个DataFrame作为表格
            tables_list = [{
                'index': 1,
                'data': df.values.tolist(),
                'shape': csv_result['shape'],
                'headers': csv_result['headers']
            }]
        else:
            text_content = f"CSV文件为空或读取失败: {path}\n"
            tables_list = []
        
        return {
            'raw_text': text_content,
            'tables': tables_list,
            'metadata': {
                'type': 'csv',
                'rows': csv_result['shape'][0],
                'columns': csv_result['shape'][1],
                'headers': csv_result['headers'],
                'encoding': csv_result.get('encoding', 'unknown')
            }
        }
    
    def extract_from_txt(path: str) -> Dict:
        """从TXT提取文本"""
        # 检测编码
        with open(path, 'rb') as f:
            raw_data = f.read(10000)  
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
        
        # 读取文件
        with open(path, 'r', encoding=encoding, errors='ignore') as f:
            text = f.read()
        
        # 获取文件大小
        file_size = Path(path).stat().st_size / 1024  # KB
        
        return {
            'raw_text': text,
            'tables': [],
            'metadata': {
                'type': 'txt',
                'encoding': encoding,
                'size_kb': round(file_size, 2),
                'characters': len(text),
                'lines': len(text.split('\n'))
            }
        }
    
    # 主流程：检测类型并调用对应函数
    file_type = detect_file_type(file_path)
        
    print(f"开始处理 {file_type} 文件: {Path(file_path).name}")
    
    if file_type == 'pdf':
        result = extract_from_pdf(file_path)
    elif file_type == 'doc':
        result = extract_from_doc(file_path)
    elif file_type == 'csv':
        result = extract_from_csv(file_path)
    else:  # txt
        result = extract_from_txt(file_path)
    
    # 添加通用元数据
    result['metadata']['file_path'] = file_path
    result['metadata']['file_name'] = Path(file_path).name
    result['metadata']['file_size'] = Path(file_path).stat().st_size
    
    print(f"完成处理 {file_type} 文件: {Path(file_path).name}")
    
    return result


def split_by_sections(text: str, keywords: Dict[str, List[str]] = None) -> Dict[str, str]:
    """
    按基础信息/指标数据/综合信息划分文本
    
    Args:
        text: 输入文本
        keywords: 自定义关键词字典
        
    Returns:
        划分后的字典
    """
    
    # 默认关键词
    default_keywords = {
        'basic': ['基础信息', '基本信息', '基本情况', '概述', '介绍', '背景', 'basic', 'overview', '公司简介'],
        'metrics': ['指标数据', '数据指标', '统计指标', '统计数据', 'metrics', 'statistics', '数据详情', '经营数据'],
        'summary': ['综合信息', '总结', '结论', '汇总', 'summary', 'conclusion', '结果', '展望', '建议']
    }
    
    # 合并用户自定义关键词
    if keywords:
        for key in default_keywords:
            if key in keywords:
                default_keywords[key].extend(keywords[key])
    
    # 初始化结果
    sections = {
        'basic': '',
        'metrics': '',
        'summary': '',
        'other': ''
    }
    
    # 按行分割文本
    lines = text.split('\n')
    current_section = 'other'
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # 检查是否为章节标题
        for section, kw_list in default_keywords.items():
            if any(kw.lower() in line_lower for kw in kw_list):
                if len(line) < 50:  # 短行更可能是标题
                    current_section = section
                    break
        
        # 将当前行添加到对应章节
        if line.strip():
            sections[current_section] += line + '\n'
    
    return sections


def split_by_rules(text: str, rules: Dict[str, List[str]] = None) -> Dict[str, str]:
    """
    使用自定义规则切分文本（备用方法）
    
    Args:
        text: 输入文本
        rules: 规则字典
        
    Returns:
        划分后的字典
    """
    
    # 默认规则
    default_rules = {
        'basic': ['公司', '企业', '名称', '地址', '时间', '日期', '负责人', '成立于', '注册'],
        'metrics': ['增长率', '百分比', '比例', '均值', '总数', '数量', '金额', '万元', '亿元', '%'],
        'summary': ['综上', '总之', '因此', '建议', '展望', '综上所述', '总体来看']
    }
    
    if rules:
        for key in default_rules:
            if key in rules:
                default_rules[key].extend(rules[key])
    
    # 按段落分割
    paragraphs = text.split('\n\n')
    sections = {'basic': '', 'metrics': '', 'summary': '', 'other': ''}
    
    for para in paragraphs:
        if not para.strip():
            continue
            
        para_lower = para.lower()
        assigned = False
        
        # 根据规则判断段落归属
        for section, kw_list in default_rules.items():
            if any(kw.lower() in para_lower for kw in kw_list):
                sections[section] += para + '\n\n'
                assigned = True
                break
        
        if not assigned:
            sections['other'] += para + '\n\n'
    
    return sections


def clean_text_data(text_dict: Dict[str, str]) -> Dict[str, str]:
    """
    清洗文本数据
    
    Args:
        text_dict: 包含各章节文本的字典
        
    Returns:
        清洗后的字典
    """
    cleaned = {}
    
    for key, text in text_dict.items():
        if not text:
            cleaned[key] = ''
            continue
        
        # 去除多余的空格和空行
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # 去除特殊字符（保留中文、英文、数字、基本标点）
        text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s\.,;:!?\-\(\)\[\]%￥元万元亿]', '', text)
        
        # 标准化日期格式
        date_patterns = [
            (r'(\d{4})年(\d{1,2})月(\d{1,2})日', r'\1-\2-\3'),
            (r'(\d{4})\.(\d{1,2})\.(\d{1,2})', r'\1-\2-\3'),
            (r'(\d{4})/(\d{1,2})/(\d{1,2})', r'\1-\2-\3'),
            (r'(\d{4})年(\d{1,2})月', r'\1-\2'),
            (r'(\d{4})\.(\d{1,2})', r'\1-\2')
        ]
        
        for pattern, repl in date_patterns:
            text = re.sub(pattern, repl, text)
        
        cleaned[key] = text.strip()
    
    return cleaned


def extract_numeric_data(text: str) -> Dict[str, List[Union[float, int]]]:
    """
    从文本中提取数值数据
    
    Args:
        text: 输入文本
        
    Returns:
        提取的数值字典
    """
    result = {
        'numbers': [],
        'percentages': [],
        'amounts': [],
        'years': []
    }
    
    if not text:
        return result
    
    # 提取所有数字
    numbers = re.findall(r'-?\d+\.?\d*', text)
    result['numbers'] = [float(n) if '.' in n else int(n) for n in numbers]
    
    # 提取百分比
    percentages = re.findall(r'(\d+\.?\d*)%', text)
    result['percentages'] = [float(p) for p in percentages]
    
    # 提取金额
    amount_patterns = [
        (r'(\d+\.?\d*)(?=\s*元)', 1),
        (r'(\d+\.?\d*)(?=\s*万元)', 10000),
        (r'(\d+\.?\d*)(?=\s*亿元)', 100000000)
    ]
    
    for pattern, multiplier in amount_patterns:
        amounts = re.findall(pattern, text)
        result['amounts'].extend([float(a) * multiplier for a in amounts])
    
    # 提取年份
    years = re.findall(r'\b(19|20)\d{2}\b', text)
    result['years'] = [int(y) for y in years]
    
    return result


def handle_missing_values(data: List, strategy: str = 'remove') -> List:
    """
    处理缺失值
    
    Args:
        data: 数据列表
        strategy: 处理策略 'remove'、'zero'、'mean'、'median'
    
    Returns:
        处理后的数据
    """
    if not data:
        return data
    
    valid_data = [x for x in data if x is not None and x != '']
    
    if strategy == 'remove':
        return valid_data
    
    elif strategy == 'zero':
        return [x if x is not None and x != '' else 0 for x in data]
    
    elif strategy == 'mean':
        numeric_data = [x for x in valid_data if isinstance(x, (int, float))]
        if numeric_data:
            mean_val = sum(numeric_data) / len(numeric_data)
            return [x if x is not None and x != '' else mean_val for x in data]
        return data
    
    elif strategy == 'median':
        numeric_data = [x for x in valid_data if isinstance(x, (int, float))]
        if numeric_data:
            sorted_data = sorted(numeric_data)
            n = len(sorted_data)
            if n % 2 == 1:
                median_val = sorted_data[n // 2]
            else:
                median_val = (sorted_data[n // 2 - 1] + sorted_data[n // 2]) / 2
            return [x if x is not None and x != '' else median_val for x in data]
        return data
    
    return data


def clean_and_extract_pipeline(text_dict: Dict[str, str], 
                               numeric_extract: bool = True,
                               missing_strategy: str = 'remove') -> Dict[str, Any]:
    """
    清洗和提取的综合流程
    
    Args:
        text_dict: 输入文本字典
        numeric_extract: 是否提取数值
        missing_strategy: 缺失值处理策略
    
    Returns:
        包含清洗后文本和提取数值的字典
    """
    # 1. 清洗文本
    cleaned_text = clean_text_data(text_dict)
    
    result = {
        'cleaned_text': cleaned_text,
        'statistics': {},
        'summary': {}
    }
    
    # 提取数值
    if numeric_extract:
        all_numbers = []
        section_stats = {}
        
        for section, text in cleaned_text.items():
            if text:
                numeric_data = extract_numeric_data(text)
                section_stats[section] = numeric_data
                
                # 收集所有数字用于整体统计
                all_numbers.extend(numeric_data['numbers'])
        
        result['statistics'] = section_stats
        
        # 整体统计
        if all_numbers:
            result['summary']['numeric_stats'] = {
                'total_numbers': len(all_numbers),
                'mean': round(sum(all_numbers) / len(all_numbers), 2),
                'max': max(all_numbers),
                'min': min(all_numbers),
                'unique_count': len(set(all_numbers))
            }
    
    # 处理缺失值
    for section in result.get('statistics', {}):
        if 'numbers' in result['statistics'][section]:
            original_count = len(result['statistics'][section]['numbers'])
            result['statistics'][section]['numbers'] = handle_missing_values(
                result['statistics'][section]['numbers'], 
                missing_strategy
            )
            result['statistics'][section]['missing_handled'] = {
                'original_count': original_count,
                'final_count': len(result['statistics'][section]['numbers']),
                'strategy': missing_strategy
            }
    
    # 添加文本统计
    result['summary']['text_stats'] = {
        'total_sections': len([v for v in cleaned_text.values() if v]),
        'total_chars': sum(len(v) for v in cleaned_text.values()),
        'sections': {k: len(v) for k, v in cleaned_text.items() if v}
    }
    
    return result


def read_large_file_in_chunks(file_path: str, chunk_size: int = 1024*1024):
    """
    分块读取大文件（生成器）
    
    Args:
        file_path: 文件路径
        chunk_size: 块大小（字节）
        
    Yields:
        文本块
    """
    # 检测编码
    with open(file_path, 'rb') as f:
        raw_data = f.read(min(10000, chunk_size))
        encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
    
    # 分块读取
    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
        while True:
            chunk = f.read(chunk_size)
            if not chunk:
                break
            yield chunk


# 使用示例
if __name__ == "__main__":
    import json
    
    # 测试文本
    sample_text = """
    基础信息
    公司名称：XX科技有限公司
    成立时间：2020年3月15日
    注册资本：1000万元
    注册地址：北京市朝阳区XX路XX号
    
    指标数据
    2023年营收增长率：25.5%
    净利润：200万元
    市场份额占比：15.3%
    员工人数：150人
    研发投入：50.8万元
    
    综合信息
    综上，公司发展态势良好，2023年各项指标均达到预期目标。
    建议继续加大研发投入，拓展市场份额。
    预计2024年营收增长率将达到30%以上。
    """
    
    print("="*50)
    print("数据预处理模块测试")
    print("="*50)
    
    # 切分测试
    print("\n1. 文本切分测试")
    sections = split_by_sections(sample_text)
    for section, content in sections.items():
        if content:
            print(f"  {section}: {len(content)} 字符")
    
    # 清洗和提取测试
    print("\n2. 数据清洗和提取测试")
    result = clean_and_extract_pipeline(sections, numeric_extract=True)
    
    # 查看提取的数值
    print("\n3. 提取的数值统计：")
    for section, stats in result['statistics'].items():
        if stats and 'numbers' in stats:
            print(f"  {section}: {stats['numbers']}")
    
    # 查看汇总统计
    print("\n4. 汇总统计：")
    print(f"  文本统计: {result['summary']['text_stats']}")
    if 'numeric_stats' in result['summary']:
        print(f"  数值统计: {result['summary']['numeric_stats']}")
    
    # 文件处理测试（如果有data文件夹）
    print("\n5. 处理data文件夹中的文件：")
    try:
        all_results = process_all_files_in_data()
        if all_results:
            print(f"  成功处理 {len(all_results)} 个文件")
            # 保存结果到JSON文件
            output_file = Path(__file__).parent / "processing_results.json"
            
            # 转换DataFrame为可序列化的格式
            serializable_results = []
            for r in all_results:
                if 'data' in r and 'tables' in r['data']:
                    # 简化表格数据以便JSON序列化
                    tables_summary = []
                    for table in r['data']['tables']:
                        if isinstance(table, dict) and 'data' in table:
                            tables_summary.append({
                                'index': table.get('index', 0),
                                'shape': table.get('shape', (0, 0)),
                                'preview': table['data'][:2] if table['data'] else []
                            })
                    r['data']['tables'] = tables_summary
                serializable_results.append(r)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(serializable_results, f, ensure_ascii=False, indent=2)
            print(f"  结果已保存到: {output_file}")
            time.sleep(10)  # 避免对文件系统的连续写入
    except Exception as e:
        print(f"  文件处理测试失败: {e}")
    
    print("程序将在10秒后退出...")
    time.sleep(10)